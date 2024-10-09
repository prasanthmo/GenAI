import os
from docx import Document

# Function to save each paragraph as a separate document in the specified directory
def split_document(input_file, output_directory):
    # Ensure the output directory exists, create it if not
    if not os.path.exists(output_directory):
        os.makedirs(output_directory)
    
    # Load the original document
    doc = Document(input_file)
    
    # Iterate through each paragraph in the original document
    for i, para in enumerate(doc.paragraphs):
        # Create a new document for each paragraph
        new_doc = Document()
        new_doc.add_paragraph(para.text)
        
        # Define the file name and path for each new document
        output_file = os.path.join(output_directory, f"Paragraph_{i+1}.docx")
        new_doc.save(output_file)
        print(f"Saved: {output_file}")

# Specify the input document path and output directory path
input_doc = os.path.abspath(r"C:\Users\mohan\GenAi\caseday1\String_Manipulation_Examples.docx")
output_dir = r'C:\Users\mohan\GenAi\caseday1'  # Use a raw string here

# Call the function to split the document and save in the output directory
split_document(input_doc, output_dir)
